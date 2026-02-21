"""
Convert Literature Review markdown files to professional academic DOCX format.
Adds student names/IDs from Pre-thesis. Saves to References DOCX folder.
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "References DOCX")
os.makedirs(OUTPUT_DIR, exist_ok=True)

STUDENT_HEADER = (
    "Md. Bokhtiar Rahman Juboraz (20301138)  |  Md. Mahir Ahnaf Ahmed (20301083)\n"
    "Department of Computer Science and Engineering, Brac University"
)

# PrimaryBlue, AccentBlue from Pre-thesis
COLOR_PRIMARY = "1A3C6E"
COLOR_ACCENT = "2563EB"
COLOR_GRAY = "64748B"


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="999999"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, COLOR_ACCENT)
        set_cell_borders(cell, COLOR_ACCENT)
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(headers)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[c_idx] if c_idx < len(row) else ""
            run = p.add_run(val)
            run.font.size = Pt(9)
            bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            set_cell_shading(cell, bg)
            set_cell_borders(cell)
    doc.add_paragraph()


def add_formatted_run(p, text, bold=False, italic=False, size=11):
    """Add formatted text handling **bold** and `code` patterns."""
    if not text:
        return
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2] + " ")
            r.bold = True
            r.font.size = Pt(size)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1] + " ")
            r.font.name = "Consolas"
            r.font.size = Pt(size - 1)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            r = p.add_run(part)
            r.font.size = Pt(size)
            if bold:
                r.bold = True
            if italic:
                r.italic = True


def process_literature_review(md_text, doc):
    """Process a literature review paper (Paper 1-17 format)."""
    lines = md_text.split("\n")
    i = 0

    # Add student header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in STUDENT_HEADER.split("\n"):
        r = p.add_run(line + "\n")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    p.paragraph_format.space_after = Pt(12)

    while i < len(lines):
        line = lines[i]

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                table_lines.append(lines[j])
                j += 1
            if len(table_lines) >= 3:
                headers = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
                rows = []
                for L in table_lines[2:]:
                    rows.append([c.strip() for c in L.strip().strip("|").split("|")])
                add_styled_table(doc, headers, rows)
            i = j
            continue

        # Main title: # Literature Review: Paper X
        if line.strip().startswith("# Literature Review:"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("#", "").strip())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            pPr = p._p.get_or_add_pPr()
            pBorders = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBorders)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # Section headers: **Abstract:**, **Introduction:**, etc.
        section_match = re.match(r"^\*\*([A-Za-z\s]+):\*\*\s*$", line.strip())
        if section_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            run = p.add_run(section_match.group(1).strip() + ":")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
            i += 1
            continue

        # Metadata lines: **Paper Title:** ... or **Authors:** ...
        meta_match = re.match(r"^\*\*([^*]+):\*\*\s*(.*)$", line.strip())
        if meta_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            run1 = p.add_run(meta_match.group(1).strip() + ": ")
            run1.bold = True
            run1.font.size = Pt(11)
            add_formatted_run(p, meta_match.group(2).strip() if meta_match.group(2) else "", size=11)
            i += 1
            continue

        # Regular paragraph (section body)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(1.25)
        add_formatted_run(p, line.strip(), size=11)
        i += 1


def process_readme(md_text, doc):
    """Process References README - different structure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in STUDENT_HEADER.split("\n"):
        r = p.add_run(line + "\n")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    p.paragraph_format.space_after = Pt(12)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                table_lines.append(lines[j])
                j += 1
            if len(table_lines) >= 3:
                headers = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
                rows = []
                for L in table_lines[2:]:
                    rows.append([c.strip() for c in L.strip().strip("|").split("|")])
                add_styled_table(doc, headers, rows)
            i = j
            continue

        if line.strip().startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("#", "").strip())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
            p.paragraph_format.space_before = Pt(14)
            i += 1
            continue

        if line.strip().startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("##", "").strip())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            p.paragraph_format.space_before = Pt(10)
            i += 1
            continue

        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBorders = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBorders)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # List items
        if line.strip().startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(2)
            add_formatted_run(p, line.strip()[2:], size=10)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        add_formatted_run(p, line.strip(), size=10)
        i += 1


def create_docx(md_path, output_name, is_readme=False):
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    if is_readme:
        process_readme(md_text, doc)
    else:
        process_literature_review(md_text, doc)

    out_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(out_path)
    print(f"Created: {output_name}")


def main():
    # All markdown files to convert
    files = [
        (os.path.join(BASE_DIR, "paper_01_adaptive_defi_borrow_lending.md"), "paper_01_adaptive_defi_borrow_lending.docx"),
        (os.path.join(BASE_DIR, "paper_02_ml_multichain_defi_fraud_detection.md"), "paper_02_ml_multichain_defi_fraud_detection.docx"),
        (os.path.join(BASE_DIR, "paper_03_risksea_ethereum_fraud.md"), "paper_03_risksea_ethereum_fraud.docx"),
        (os.path.join(BASE_DIR, "paper_04_lime_shap_loan_approval.md"), "paper_04_lime_shap_loan_approval.docx"),
        (os.path.join(BASE_DIR, "paper_05_rl_interest_rate_defi_lending.md"), "paper_05_rl_interest_rate_defi_lending.docx"),
        (os.path.join(BASE_DIR, "paper_06_rl_credit_scoring_underwriting.md"), "paper_06_rl_credit_scoring_underwriting.docx"),
        (os.path.join(BASE_DIR, "paper_07_cbdc_financial_inclusion.md"), "paper_07_cbdc_financial_inclusion.docx"),
        (os.path.join(BASE_DIR, "paper_08_blockchain_financial_inclusion_sustainable_dev.md"), "paper_08_blockchain_financial_inclusion_sustainable_dev.docx"),
        (os.path.join(BASE_DIR, "paper_09_defi_lending_evaluation_system.md"), "paper_09_defi_lending_evaluation_system.docx"),
        (os.path.join(BASE_DIR, "paper_10_sok_defi", "paper_10_sok_defi.md"), "paper_10_sok_defi.docx"),
        (os.path.join(BASE_DIR, "paper_11_shap_unified_approach", "paper_11_shap_unified_approach.md"), "paper_11_shap_unified_approach.docx"),
        (os.path.join(BASE_DIR, "paper_12_isolation_forest.md"), "paper_12_isolation_forest.docx"),
        (os.path.join(BASE_DIR, "paper_13_ethereum_smart_contract_attacks_sok.md"), "paper_13_ethereum_smart_contract_attacks_sok.docx"),
        (os.path.join(BASE_DIR, "paper_14_governance_blockchain_economy.md"), "paper_14_governance_blockchain_economy.docx"),
        (os.path.join(BASE_DIR, "paper_15_smart_contract_formal_verification.md"), "paper_15_smart_contract_formal_verification.docx"),
        (os.path.join(BASE_DIR, "paper_16_bitcoin_ml_security.md"), "paper_16_bitcoin_ml_security.docx"),
        (os.path.join(BASE_DIR, "paper_17_ml_explainability_finance_default_risk", "paper_17_ml_explainability_finance_default_risk.md"), "paper_17_ml_explainability_finance_default_risk.docx"),
        (os.path.join(BASE_DIR, "References", "README.md"), "References_README.docx"),
    ]

    for md_path, out_name in files:
        if os.path.exists(md_path):
            create_docx(md_path, out_name, is_readme="README" in out_name)
        else:
            print(f"Skipped (not found): {md_path}")

    print("\nAll DOCX files generated in References DOCX folder.")


if __name__ == "__main__":
    main()
