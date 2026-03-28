"""Convert PROJECT_REPORT.md to PROJECT_REPORT.docx using python-docx."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = Path(__file__).parent.parent / "Documentation" / "PROJECT_REPORT.md"
DOCX_PATH = MD_PATH.with_suffix(".docx")


def set_heading_style(paragraph, level: int):
    """Apply colour and font to heading paragraphs."""
    colour = RGBColor(0x1A, 0x3C, 0x6E)  # PrimaryBlue
    for run in paragraph.runs:
        run.font.color.rgb = colour
        run.font.bold = True
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11)


def add_horizontal_rule(doc):
    """Add a thin horizontal line paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2563EB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_inline(run_text: str, paragraph):
    """Handle **bold**, *italic*, and `code` inline markdown in a paragraph."""
    # Split on bold, italic, code markers
    token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = token_re.split(run_text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def set_body_para(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = Pt(13)


def build_docx(md_text: str) -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # --- Skip mermaid/code blocks ---
        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        # --- Horizontal rule ---
        if re.match(r"^---+$", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # --- Headings ---
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", m.group(2))  # strip links
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)                # strip bold/italic
            word_level = min(level, 4)
            para = doc.add_heading(text, level=word_level)
            set_heading_style(para, level)
            para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
            para.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # --- Bullet list item ---
        if re.match(r"^(\s*)[-*]\s+", line):
            indent = len(re.match(r"^(\s*)", line).group(1))
            text = re.sub(r"^\s*[-*]\s+", "", line)
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            para = doc.add_paragraph(style="List Bullet")
            if indent > 0:
                para.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
            para.paragraph_format.space_after = Pt(3)
            parse_inline(text, para)
            i += 1
            continue

        # --- Numbered list item ---
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_after = Pt(3)
            parse_inline(text, para)
            i += 1
            continue

        # --- Table row ---
        if line.strip().startswith("|"):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # Filter out separator rows (---|---)
            data_rows = [r for r in table_lines if not re.match(r"^\s*\|[\s\-:|]+\|\s*$", r)]
            if not data_rows:
                continue

            def parse_row(row_str):
                cells = [c.strip() for c in row_str.strip().strip("|").split("|")]
                return cells

            rows = [parse_row(r) for r in data_rows]
            col_count = max(len(r) for r in rows)

            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    if ci >= col_count:
                        break
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cell_text)
                    clean = re.sub(r"\*+([^*]+)\*+", r"\1", clean)
                    clean = re.sub(r"`([^`]+)`", r"\1", clean)
                    p = cell.paragraphs[0]
                    run = p.add_run(clean)
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
                        cell._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), "DBEAFE")
                        cell._tc.tcPr.append(shd)
            doc.add_paragraph()  # spacing after table
            continue

        # --- Bold metadata lines (key: value at top of doc) ---
        if re.match(r"^\*\*[^*]+\*\*", line) and ":" in line:
            para = doc.add_paragraph()
            set_body_para(para)
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
            parse_inline(text, para)
            i += 1
            continue

        # --- Empty line ---
        if not line.strip():
            i += 1
            continue

        # --- Regular paragraph ---
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        if text.strip():
            para = doc.add_paragraph()
            set_body_para(para)
            parse_inline(text, para)
        i += 1

    return doc


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = build_docx(md_text)
    doc.save(str(DOCX_PATH))
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
